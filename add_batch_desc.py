
"""
Fixed script to add batch description with new filename to avoid PermissionError.
"""
import os
import sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout = open('d:/web-lcd/add_batch_desc_log.txt', 'w', encoding='utf-8')
sys.stderr = sys.stdout

# Look for the latest "Final" file
input_path = 'd:/web-lcd/Bao cao KLTN - Final.docx'
output_path = 'd:/web-lcd/Bao cao KLTN - Final Updated.docx'

print(f"Opening: {input_path}")
doc = Document(input_path)

def insert_paragraph_after(ref_para, text, style_name='Normal'):
    new_p_elem = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), style_name)
    pPr.append(pStyle)
    new_p_elem.append(pPr)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    new_p_elem.append(r)
    ref_para._p.addnext(new_p_elem)
    for p in doc.paragraphs:
        if p._element is new_p_elem:
            return p
    return None

target_fig = "Hình 4.16"
idx = -1
for i, p in enumerate(doc.paragraphs):
    if target_fig in p.text and p.style.name == 'Caption' and i > 400:
        idx = i
        print(f"Found {target_fig} at index {i}")
        break

if idx != -1:
    desc = (
        "Chức năng cập nhật thông tin thành viên hàng loạt là một công cụ mạnh mẽ giúp quản trị viên tối ưu hóa "
        "thời gian khi cần xử lý dữ liệu cho nhiều tài khoản cùng lúc. Quản trị viên có thể lựa chọn danh sách "
        "thành viên cần xử lý thông qua các ô checkbox tích hợp sẵn trong bảng danh sách. Sau khi chọn, "
        "hệ thống sẽ hiển thị giao diện cập nhật tập trung cho phép thay đổi đồng loạt các thông tin như: "
        "đơn vị công tác (Ban chuyên môn), chức vụ, khóa học, lớp sinh hoạt hoặc trạng thái kích hoạt tài khoản. "
        "Cơ chế này đặc biệt hữu ích trong các đợt kiện toàn nhân sự hoặc chuyển giao nhiệm kỳ của Liên Chi đoàn, "
        "giúp đảm bảo tính đồng bộ và chính xác tuyệt đối của dữ liệu thành viên mà không cần thao tác thủ công "
        "trên từng hồ sơ riêng lẻ."
    )
    insert_paragraph_after(doc.paragraphs[idx], desc)
    doc.save(output_path)
    print(f"Successfully added description after {target_fig} and saved to {output_path}")
else:
    print(f"Could not find {target_fig}")

sys.stdout.close()
